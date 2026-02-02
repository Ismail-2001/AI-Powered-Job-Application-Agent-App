"""
Document Builder Utility
Role: Generate professional, ATS-friendly DOCX files.
"""

from typing import Dict, Any, List
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class DocumentBuilder:
    """
    Handles creation and formatting of MS Word documents.
    """

    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Configure document styles for ATS readability"""
        # Set margins (standard 1 inch)
        for section in self.doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        # Standard font
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

    def create_cv(self, cv_data: Dict[str, Any], output_path: str):
        """
        Generate a CV document from structured data.

        Args:
            cv_data: Dictionary containing 'personal_info', 'experience', 'education', 'skills'
            output_path: File path to save the DOCX
        """
        try:
            # 1. Header (Name & Contact)
            self._add_header(cv_data.get('personal_info', {}))

            # 2. Professional Summary
            if 'summary' in cv_data:
                self._add_section_title("PROFESSIONAL SUMMARY")
                self.doc.add_paragraph(cv_data['summary'])

            # 3. Skills
            if 'skills' in cv_data:
                self._add_section_title("CORE SKILLS")
                self._add_skills(cv_data['skills'])

            # 4. Experience
            if 'experience' in cv_data:
                self._add_section_title("PROFESSIONAL EXPERIENCE")
                for role in cv_data['experience']:
                    self._add_experience_item(role)

            # 5. Education
            if 'education' in cv_data:
                self._add_section_title("EDUCATION")
                for edu in cv_data['education']:
                    self._add_education_item(edu)
            
            # Save
            self.doc.save(output_path)
            print(f"✅ Document saved to: {output_path}")

        except Exception as e:
            print(f"❌ Failed to create document: {e}")
            raise

    def _add_header(self, info: Dict[str, str]):
        """Add personal info header"""
        name = info.get('name', 'Candidate Name')
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = p.add_run(name)
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0, 0, 0) # Black

        # Contact line
        contact_parts = []
        if info.get('email'): contact_parts.append(info['email'])
        if info.get('phone'): contact_parts.append(info['phone'])
        if info.get('linkedin'): contact_parts.append(info['linkedin'])
        if info.get('location'): contact_parts.append(info['location'])
        
        if contact_parts:
            p = self.doc.add_paragraph(" | ".join(contact_parts))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(10)

    def _add_section_title(self, title: str):
        """Add a standardized section header"""
        p = self.doc.add_paragraph()
        p.space_before = Pt(12)
        p.space_after = Pt(6)
        
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Calibri'
        
        # Add bottom border style (hacky in python-docx, usually simple underline is safer for ATS)
        # Using simple underline for safety
        run.underline = True

    def _add_skills(self, skills: Any):
        """Format skills section"""
        if isinstance(skills, list):
            # If simple list, join with bullets or pipes? 
            # ATS prefers comma separated or bullet points.
            self.doc.add_paragraph(", ".join(skills))
        elif isinstance(skills, dict):
            # Categorized skills
            for category, items in skills.items():
                p = self.doc.add_paragraph()
                run = p.add_run(f"{category}: ")
                run.bold = True
                p.add_run(", ".join(items))

    def _add_experience_item(self, role: Dict[str, Any]):
        """Add a job role"""
        # Title line
        p = self.doc.add_paragraph()
        p.space_before = Pt(8)
        
        # Company Name (Bold)
        company = role.get('company', '')
        r1 = p.add_run(company)
        r1.bold = True
        
        # Location (Right align if possible, but keep simple for ATS)
        location = role.get('location', '')
        if location:
            p.add_run(f" — {location}")

        # Job Title & Dates
        p2 = self.doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        title = role.get('title', '')
        dates = role.get('dates', '')
        
        r2 = p2.add_run(title)
        r2.italic = True
        if dates:
            p2.add_run(f" | {dates}")

        # Bullets
        achievements = role.get('achievements', role.get('responsibilities', []))
        for item in achievements:
            self.doc.add_paragraph(item, style='List Bullet')

    def _add_education_item(self, edu: Dict[str, Any]):
        """Add education item"""
        p = self.doc.add_paragraph()
        p.space_before = Pt(6)
        
        school = edu.get('school', '')
        degree = edu.get('degree', '')
        dates = edu.get('dates', '')
        
        r = p.add_run(school)
        r.bold = True
        p.add_run(f" — {degree}")
        if dates:
            p.add_run(f" ({dates})")

    def create_cover_letter(self, letter_body: str, profile: Dict[str, Any], output_path: str):
        """
        Generate a Cover Letter document.
        
        Args:
            letter_body: The text content of the letter
            profile: Candidate profile (for header)
            output_path: File path to save
        """
        try:
            # Re-initialize doc for new file
            self.doc = Document()
            self._setup_styles()
            
            # 1. Header (Same as CV)
            self._add_header(profile.get('personal_info', {}))
            
            # 2. Spacing
            self.doc.add_paragraph().space_after = Pt(24)
            
            # 3. Body
            # Split by newlines to create proper paragraphs
            for paragraph in letter_body.split('\n'):
                if paragraph.strip():
                    p = self.doc.add_paragraph(paragraph.strip())
                    p.paragraph_format.space_after = Pt(12)
            
            # Save
            self.doc.save(output_path)
            print(f"✅ Cover Letter saved to: {output_path}")
            
        except Exception as e:
            print(f"❌ Failed to create cover letter: {e}")
            raise
